from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pandas as pd

def insert_text_at_bookmark(doc: Document, bookmark: str, text: str) -> bool:
    """
    Reemplaza la cadena 'bookmark' (p.ej. {{RESUMEN}}) por 'text' en el documento.
    - Arial 11, sin negrita.
    - Devuelve True si reemplazó algo.
    """
    replaced = False
    for p in doc.paragraphs:
        if bookmark in p.text:
            before, _, after = p.text.partition(bookmark)
            p.text = ""  # limpiar párrafo
            r1 = p.add_run(before)
            r2 = p.add_run(text)
            r3 = p.add_run(after)
            for r in (r1, r2, r3):
                r.font.name = "Arial"
                r.font.size = Pt(11)
            replaced = True
    return replaced

def _set_table_borders(table):
    """Aplica bordes finos negros a toda la tabla si no hay estilo."""
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '8')         # grosor ≈ 0.5pt
        el.set(qn('w:color'), '000000') # negro
        borders.append(el)
    tblPr.append(borders)
    tbl._element.append(tblPr)

def _shade_cell(cell, hex_fill):
    """Aplica color de fondo a una celda."""
    tc_pr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_fill)
    tc_pr.append(shd)

def insert_dataframe_at_bookmark(doc: Document, df: pd.DataFrame, bookmark: str) -> bool:
    """
    Inserta una tabla en el marcador con formato institucional:
    - Encabezado azul (1B355E) con texto blanco, negrita Arial 8.
    - Cuerpo en Arial 8, negro.
    - Bordes negros finos si no hay estilo de tabla.
    """
    if df is None or df.empty:
        return False

    temp = df.copy()
    if temp.index.name or not isinstance(temp.index, pd.RangeIndex):
        idx_name = temp.index.name or " "
        temp = temp.reset_index().rename(columns={temp.columns[0]: idx_name})

    temp = temp.fillna("")
    temp.columns = temp.columns.astype(str)

    for p in doc.paragraphs:
        if bookmark in p.text:
            # Crear tabla
            table = doc.add_table(rows=1, cols=len(temp.columns))

            # Intentar aplicar estilo existente
            for style_name in ("Table Grid", "TableGrid", "Normal Table", "Tabla con cuadrícula"):
                try:
                    table.style = style_name
                    break
                except Exception:
                    pass
            else:
                _set_table_borders(table)

            # === ENCABEZADO ===
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(temp.columns):
                hdr_cells[i].text = str(col)
                run = hdr_cells[i].paragraphs[0].runs[0]
                run.font.name = "Arial"
                run.font.size = Pt(8)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # blanco
                _shade_cell(hdr_cells[i], "1B355E")           # azul institucional

            # === CUERPO ===
            for _, row in temp.iterrows():
                row_cells = table.add_row().cells
                for i, val in enumerate(row):
                    row_cells[i].text = str(val)
                    run = row_cells[i].paragraphs[0].runs[0]
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(0, 0, 0)

            # Reemplazar marcador
            before, _, after = p.text.partition(bookmark)
            p.text = before
            anchor = p.add_run()
            anchor._r.addnext(table._element)
            if after:
                tail = p.add_run(after)
                tail.font.name = "Arial"
                tail.font.size = Pt(8)

            return True
    return False